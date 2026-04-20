from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree
from latex2mathml.converter import convert as latex_to_mathml


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "דף נוסחאות משולב.docx"
MML2OMML = Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")


if not MML2OMML.exists():
    raise FileNotFoundError(f"Missing Word math transform: {MML2OMML}")


MATH_TRANSFORM = etree.XSLT(etree.parse(str(MML2OMML)))


def set_paragraph_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    if not p_pr.xpath("./w:bidi"):
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        p_pr.append(bidi)


def set_run_rtl(run):
    r_pr = run._r.get_or_add_rPr()
    if not r_pr.xpath("./w:rtl"):
        rtl = OxmlElement("w:rtl")
        rtl.set(qn("w:val"), "1")
        r_pr.append(rtl)


def add_text(doc, text, *, style="Normal", bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4):
    p = doc.add_paragraph(style=style)
    set_paragraph_rtl(p, align=align)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_rtl(run)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_rtl(p)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_rtl(run)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_rtl(p)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_rtl(run)
    return p


def add_equation(doc, latex, note=None):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(3)
    mathml = latex_to_mathml(latex)
    omml = MATH_TRANSFORM(etree.fromstring(mathml.encode("utf-8"))).getroot()
    p._element.append(omml)
    if note:
        add_text(doc, note, italic=True, space_after=4)


def add_note(doc, label, text):
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{label}: ")
    set_run_rtl(r1)
    r1.bold = True
    r2 = p.add_run(text)
    set_run_rtl(r2)
    return p


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)

    for style_name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True

    if "Compact Note" not in doc.styles:
        style = doc.styles.add_style("Compact Note", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Aptos"
        style.font.size = Pt(10)


def build_document():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    set_paragraph_rtl(title, align=WD_ALIGN_PARAGRAPH.CENTER)
    title_run = title.add_run("דף נוסחאות משולב")
    set_run_rtl(title_run)

    subtitle = doc.add_paragraph()
    set_paragraph_rtl(subtitle, align=WD_ALIGN_PARAGRAPH.CENTER)
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("הסתברות, תהליכים אקראיים, שערוך ומרקוב")
    set_run_rtl(run)
    run.italic = True

    add_text(
        doc,
        "המסמך מאחד את שלושת הדפים לקובץ לימוד אחד: נוסחאות מרכזיות, תנאים לשימוש, והבדלים שקל להתבלבל בהם במבחן.",
        space_after=6,
    )
    add_note(
        doc,
        "סימונים",
        "E הוא תוחלת, Var שונות, Cov קווריאנס, R אוטוקורלציה/קרוס-קורלציה, S צפיפות הספק, P מטריצת מעבר.",
    )

    add_heading(doc, "1. יסודות הסתברות ותוחלת מותנית", level=1)
    add_equation(doc, r"E[aX+bY+c]=aE[X]+bE[Y]+c")
    add_equation(doc, r"\mathrm{Var}(X)=E[X^2]-(E[X])^2")
    add_equation(doc, r"\mathrm{Cov}(X,Y)=E[XY]-E[X]E[Y]")
    add_bullet(doc, "אם X ו-Y בלתי תלויים אז Cov(X,Y)=0, אבל הכיוון ההפוך בדרך כלל לא נכון.")
    add_equation(doc, r"E[X]=E[E[X\mid Y]]")
    add_equation(doc, r"\mathrm{Var}(X)=E[\mathrm{Var}(X\mid Y)]+\mathrm{Var}(E[X\mid Y])")
    add_equation(doc, r"P(A)=\sum_i P(A\mid B_i)P(B_i)")
    add_equation(doc, r"P(A)=\int P(A\mid Y=y)f_Y(y)\,dy")
    add_equation(doc, r"P(B_j\mid A)=\frac{P(A\mid B_j)P(B_j)}{\sum_i P(A\mid B_i)P(B_i)}")
    add_equation(doc, r"f_{X\mid Y}(x\mid y)=\frac{f_{X,Y}(x,y)}{f_Y(y)}")
    add_equation(doc, r"f_{Y\mid X}(y\mid x)=\frac{f_{X\mid Y}(x\mid y)f_X(x)}{f_Y(y)}")
    add_note(doc, "מתי להשתמש", "כשיש משתנה נסתר שמקל את החישוב, מתנים עליו. כשמבקשים להפוך 'תוצאה -> סיבה', כמעט תמיד בייס.")
    add_note(doc, "טיפ קלאסי", "בסכום מקרי או בתהליך שמתבצע בשלבים, הסתברות שלמה ותוחלת שלמה הן בדרך כלל נקודת הפתיחה.")

    add_heading(doc, "2. התפלגויות חשובות וזהויות בסיס", level=1)
    add_heading(doc, "התפלגויות בדידות", level=2)
    add_equation(doc, r"X\sim\mathrm{Bernoulli}(p),\quad P(X=1)=p,\quad E[X]=p,\quad \mathrm{Var}(X)=p(1-p)")
    add_equation(doc, r"X\sim\mathrm{Bin}(n,p),\quad P(X=k)=\binom{n}{k}p^k(1-p)^{n-k}")
    add_equation(doc, r"E[X]=np,\quad \mathrm{Var}(X)=np(1-p)")
    add_equation(doc, r"X\sim\mathrm{Pois}(\lambda),\quad P(X=k)=e^{-\lambda}\frac{\lambda^k}{k!}")
    add_equation(doc, r"E[X]=\lambda,\quad \mathrm{Var}(X)=\lambda")
    add_equation(doc, r"K\sim\mathrm{Geom}(p),\quad P(K=k)=(1-p)^k p,\quad k=0,1,2,\dots")
    add_equation(doc, r"E[K]=\frac{1-p}{p},\quad \mathrm{Var}(K)=\frac{1-p}{p^2}")
    add_bullet(doc, "אם הקורס מגדיר גיאומטרית כמספר הניסיונות עד הצלחה, אז T=K+1 ולכן E[T]=1/p.")
    add_heading(doc, "התפלגויות רציפות", level=2)
    add_equation(doc, r"X\sim U[a,b],\quad f_X(x)=\frac{1}{b-a},\ a\le x\le b")
    add_equation(doc, r"E[X]=\frac{a+b}{2},\quad \mathrm{Var}(X)=\frac{(b-a)^2}{12}")
    add_equation(doc, r"T\sim\mathrm{Exp}(\lambda),\quad f_T(t)=\lambda e^{-\lambda t},\ t\ge 0")
    add_equation(doc, r"E[T]=\frac{1}{\lambda},\quad \mathrm{Var}(T)=\frac{1}{\lambda^2}")
    add_equation(doc, r"P(T>s+t\mid T>s)=P(T>t)")
    add_equation(doc, r"X\sim N(\mu,\sigma^2),\quad f_X(x)=\frac{1}{\sqrt{2\pi\sigma^2}}e^{-\frac{(x-\mu)^2}{2\sigma^2}}")
    add_equation(doc, r"\phi_X(\omega)=\exp\left(j\mu\omega-\frac{1}{2}\sigma^2\omega^2\right)")
    add_equation(doc, r"X\sim N(0,\sigma^2)\ \Longrightarrow\ E[X^4]=3\sigma^4")
    add_bullet(doc, "אם X גאוסי אז |X| בדרך כלל אינו גאוסי.")
    add_heading(doc, "זהויות מהירות", level=2)
    add_equation(doc, r"e^{j\theta}=\cos\theta+j\sin\theta")
    add_equation(doc, r"\cos\theta=\frac{e^{j\theta}+e^{-j\theta}}{2},\qquad \sin\theta=\frac{e^{j\theta}-e^{-j\theta}}{2j}")
    add_equation(doc, r"\operatorname{sinc}(x)=\frac{\sin(\pi x)}{\pi x}")
    add_equation(doc, r"\delta(t-t_0)\ \longleftrightarrow\ e^{-j2\pi f t_0}")
    add_bullet(doc, "אינטגרל של פונקציה אי-זוגית על תחום סימטרי סביב 0 שווה ל-0.")

    add_heading(doc, "3. וקטורים גאוסיים וגאוסיות במשותף", level=1)
    add_equation(doc, r"a^T X\ \mathrm{is\ Gaussian\ for\ every}\ a\ \Longleftrightarrow\ X\ \mathrm{jointly\ Gaussian}")
    add_equation(doc, r"\phi_X(\omega)=\exp\left(j\omega^T m_X-\frac{1}{2}\omega^T C_X \omega\right)")
    add_equation(doc, r"Y=AX+b,\ X\ \mathrm{Gaussian} \Longrightarrow Y\ \mathrm{Gaussian}")
    add_equation(doc, r"m_Y=Am_X+b,\qquad C_Y=AC_XA^T")
    add_bullet(doc, "כל תת-וקטור של וקטור גאוסי במשותף הוא גם גאוסי במשותף.")
    add_bullet(doc, "גאוסיות של כל רכיב לחוד אינה מספיקה לגאוסיות במשותף. צריך לבדוק כל צירוף ליניארי.")
    add_equation(doc, r"\mathrm{Cov}(X,Y)=0,\ X,Y\ \mathrm{jointly\ Gaussian} \Longrightarrow X\perp Y")
    add_bullet(doc, "באופן כללי: אי-תלות גוררת חוסר קורלציה, אבל חוסר קורלציה לא גורר אי-תלות.")
    add_equation(doc, r"\hat X_{\mathrm{MMSE}}(Y)=E[X\mid Y]")
    add_equation(doc, r"\hat X_{\mathrm{LMMSE}}=m_X+C_{XY}C_Y^{-1}(Y-m_Y)")
    add_bullet(doc, "אם X ו-Y גאוסים במשותף, אז MMSE = LMMSE והאומדן האופטימלי הוא ליניארי.")
    add_equation(doc, r"E[X\mid Y=y]=m_X+C_{XY}C_{YY}^{-1}(y-m_Y)")
    add_equation(doc, r"\mathrm{Cov}(X\mid Y)=C_{XX}-C_{XY}C_{YY}^{-1}C_{YX}")
    add_note(doc, "מתי להשתמש", "כשרואים וקטור גאוסי או תצפית ליניארית עם רעש גאוסי, נוסחת ההתניה של גאוסי היא בדרך כלל הפתרון המהיר.")

    add_heading(doc, "4. תהליך פואסון וזמני המתנה", level=1)
    add_bullet(doc, "תהליך פואסון הומוגני מוגדר על ידי N(0)=0, תוספות בלתי תלויות, ותוספות סטציונריות.")
    add_equation(doc, r"N(t)-N(s)\sim\mathrm{Pois}(\lambda(t-s)),\qquad 0\le s<t")
    add_equation(doc, r"P(N(t)=k)=e^{-\lambda t}\frac{(\lambda t)^k}{k!}")
    add_equation(doc, r"E[N(t)]=\lambda t,\qquad \mathrm{Var}(N(t))=\lambda t")
    add_equation(doc, r"E[N(t)N(s)]=\lambda\min(t,s)+\lambda^2 ts")
    add_equation(doc, r"T_1\sim\mathrm{Exp}(\lambda)")
    add_equation(doc, r"S_n=T_1+\cdots+T_n\sim\mathrm{Erlang}(n,\lambda)")
    add_bullet(doc, "זמני ההמתנה בין הגעות עוקבות הם i.i.d אקספוננציאליים עם אותו קצב λ.")
    add_bullet(doc, "בקטעים זרים, מספרי ההגעות בלתי תלויים.")
    add_bullet(doc, "אם מופיע פרמטר אקראי בתוך קצב פואסון, התניה על הפרמטר ואז החלקה/תוחלת שלמה היא הדרך הנכונה.")
    add_bullet(doc, "סכום של תהליכים פואסוניים יכול להיות פואסון בתנאים מתאימים; מכפלה של תהליכים פואסוניים בדרך כלל אינה תהליך פואסון.")

    add_heading(doc, "5. סטציונריות, סמ\"ר, אוטוקורלציה ו-PSD", level=1)
    add_equation(doc, r"(X(t_1),\dots,X(t_n)) \stackrel{d}{=} (X(t_1+\tau),\dots,X(t_n+\tau))")
    add_equation(doc, r"E[X(t)]=m_X,\qquad R_X(\tau)=E[X(t+\tau)X^*(t)]")
    add_equation(doc, r"R_{XY}(\tau)=E[X(t+\tau)Y^*(t)]")
    add_equation(doc, r"E[W(t)]=0,\qquad R_W(\tau)=\sigma_W^2\delta(\tau)")
    add_equation(doc, r"S_X(f)=\mathcal{F}\{R_X(\tau)\},\qquad R_X(\tau)=\int_{-\infty}^{\infty} S_X(f)e^{j2\pi f\tau}\,df")
    add_equation(doc, r"P_X=R_X(0)=\int_{-\infty}^{\infty} S_X(f)\,df")
    add_equation(doc, r"Y=h*X,\ X\ \mathrm{WSS} \Longrightarrow S_Y(f)=|H(f)|^2S_X(f)")
    add_equation(doc, r"S_{YX}(f)=H(f)S_X(f),\qquad S_{XY}(f)=H^*(f)S_X(f)")
    add_bullet(doc, "הנוסחאות לקרוס-ספקטרום נכתבו לפי ההגדרה R_XY(τ)=E[X(t+τ)Y*(t)].")
    add_bullet(doc, "סטציונרי חזק דורש אי-תלות בזמן של כל ההתפלגויות; סמ\"ר דורש רק תוחלת קבועה וקורלציה שתלויה בהפרש זמנים.")
    add_bullet(doc, "תהליך עם תנאי עצירה, התחלה מיוחדת או כלל שמתעדכן עם הזמן הוא חשוד מיידית כלא-סטציונרי.")

    add_heading(doc, "6. שערוך MMSE, שערוך ליניארי ומסנן וינר", level=1)
    add_equation(doc, r"\hat X_{\mathrm{MMSE}}(Y)=E[X\mid Y]")
    add_equation(doc, r"E[(X-\hat X_{\mathrm{MMSE}})g(Y)]=0")
    add_equation(doc, r"\hat X=aY+b,\qquad a=\frac{\mathrm{Cov}(X,Y)}{\mathrm{Var}(Y)},\qquad b=E[X]-aE[Y]")
    add_equation(doc, r"\mathrm{MSE}_{\min}=\mathrm{Var}(X)-\frac{\mathrm{Cov}(X,Y)^2}{\mathrm{Var}(Y)}")
    add_equation(doc, r"\hat X=m_X+C_{XY}C_Y^{-1}(Y-m_Y)")
    add_equation(doc, r"H_{\mathrm{opt}}(f)=\frac{S_{XY}(f)}{S_Y(f)}")
    add_equation(doc, r"S_e(f)=S_X(f)-\frac{|S_{XY}(f)|^2}{S_Y(f)}")
    add_equation(doc, r"\mathrm{MSE}=\int_{-\infty}^{\infty} S_e(f)\,df")
    add_equation(doc, r"Y=X+N,\ X\perp N \Longrightarrow H_{\mathrm{opt}}(f)=\frac{S_X(f)}{S_X(f)+S_N(f)}")
    add_bullet(doc, "במסנן וינר עובדים תמיד באותו סדר: אוטוקורלציה/קרוס-קורלציה -> PSD -> הצבת H_opt.")
    add_bullet(doc, "אם X ו-Y חסרי קורלציה, המסנן הליניארי האופטימלי הוא 0, והאומדן הקבוע הטוב ביותר הוא E[X].")
    add_bullet(doc, "אם המשתנים גאוסים במשותף, ה-MMSE כבר ליניארי; אם לא, MMSE עלול להיות לא ליניארי וטוב יותר מכל מסנן ליניארי.")

    add_heading(doc, "7. שרשראות מרקוב", level=1)
    add_equation(doc, r"P(X_{n+1}=j\mid X_n=i,\ldots,X_0)=P_{ij}")
    add_equation(doc, r"P^{(n)}_{ij}=(P^n)_{ij}")
    add_equation(doc, r"\mu^{(n)}=\mu^{(0)}P^n")
    add_equation(doc, r"\pi=\pi P,\qquad \sum_i \pi_i=1")
    add_equation(doc, r"P(X_{2n}=k\mid X_n=j)=(P^n)_{jk}")
    add_equation(doc, r"P(X_0=i\mid X_n=j)=\frac{\mu_i^{(0)}(P^n)_{ij}}{\mu_j^{(n)}}")
    add_equation(doc, r"E_j[N_i]=\sum_{n=0}^{\infty}(P^n)_{ji}")
    add_bullet(doc, "הומוגנית בזמן = חוקי המעבר לא משתנים עם n. סטציונרית = ההתפלגות על המצבים לא משתנה אחרי הפעלת P.")
    add_bullet(doc, "במחלקה קשירה וסגורה המצבים נשנים; במחלקה פתוחה המצבים חולפים.")
    add_bullet(doc, "בשרשרת סופית ואי-פריקה קיים פילוג סטציונרי יחיד, והוא חיובי על כל המצבים.")
    add_bullet(doc, "אם התהליך i.i.d אז הוא מרקובי הומוגני, אבל לא כל תהליך מרקובי הוא i.i.d.")

    add_heading(doc, "8. מלכודות נפוצות למבחן", level=1)
    add_bullet(doc, "חוסר קורלציה אינו אי-תלות, אלא אם המשתנים גאוסים במשותף.")
    add_bullet(doc, "כל רכיב גאוסי לחוד עדיין לא מבטיח שהווקטור גאוסי במשותף.")
    add_bullet(doc, "סטציונרי חזק חזק יותר מסמ\"ר; סמ\"ר מדבר רק על רגע ראשון ושני.")
    add_bullet(doc, "MMSE הוא תמיד E[X|Y]. מסנן וינר הוא הפתרון רק תחת אילוץ ליניאריות או תחת גאוסיות במשותף.")
    add_bullet(doc, "הומוגניות במרקוב היא תכונה של P; פילוג סטציונרי הוא תכונה של התפלגות π.")
    add_bullet(doc, "כשרואים sign(X) או |X|, נסו לחזור ל-X המקורי דרך סימטריה או חלוקה למקרים.")
    add_bullet(doc, "בבדיקת אי-תלות של משתנה 'בוחר', שאלו אם המקורות שהוא בוחר ביניהם באמת זהים בהתפלגות שלהם.")

    add_text(
        doc,
        "הערה מסכמת: זהו דף עזר ללמידה ולפתרון מהיר. אם בשאלה חסר צעד, התחילו תמיד מהגדרה, בדקו תנאים, ורק אחר כך בחרו את הנוסחה.",
        bold=True,
        space_after=0,
    )

    doc.save(str(OUTPUT))


if __name__ == "__main__":
    build_document()
    print(OUTPUT.as_posix().encode("ascii", "backslashreplace").decode("ascii"))
