from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree
from latex2mathml.converter import convert as latex_to_mathml


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "דף נוסחאות משולב - גרסה 2.docx"
MML2OMML = Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")


if not MML2OMML.exists():
    raise FileNotFoundError(f"Missing Word math transform: {MML2OMML}")


MATH_TRANSFORM = etree.XSLT(etree.parse(str(MML2OMML)))


def set_paragraph_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    if not p_pr.xpath("./w:bidi"):
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        p_pr.append(bidi)


def set_run_rtl(run):
    r_pr = run._r.get_or_add_rPr()
    if not r_pr.xpath("./w:rtl"):
        rtl = OxmlElement("w:rtl")
        rtl.set(qn("w:val"), "1")
        r_pr.append(rtl)


def add_paragraph(doc, text="", style="Normal", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    set_paragraph_rtl(p, align=align)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_rtl(run)
        run.bold = bold
        run.italic = italic
    return p


def add_heading(doc, text, level=1):
    return add_paragraph(doc, text, style=f"Heading {level}", space_after=3, bold=True)


def add_bullet(doc, text):
    return add_paragraph(doc, text, style="List Bullet", space_after=2)


def add_note(doc, label, text):
    p = add_paragraph(doc, "", space_after=3)
    r1 = p.add_run(f"{label}: ")
    set_run_rtl(r1)
    r1.bold = True
    r2 = p.add_run(text)
    set_run_rtl(r2)
    return p


def add_equation(doc, latex):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(3)
    mathml = latex_to_mathml(latex)
    omml = MATH_TRANSFORM(etree.fromstring(mathml.encode("utf-8"))).getroot()
    p._element.append(omml)
    return p


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    for name, size, bold in [("Normal", 11, False), ("Title", 18, True), ("Heading 1", 15, True), ("Heading 2", 12, True)]:
        style = doc.styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = bold

    if "SmallNote" not in doc.styles:
        style = doc.styles.add_style("SmallNote", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Aptos"
        style.font.size = Pt(10)


def add_distribution_block(doc, name, use, formula, mean, var):
    add_paragraph(doc, name, style="Heading 2", space_after=2)
    add_note(doc, "מתי משתמשים", use)
    add_equation(doc, formula)
    add_equation(doc, mean)
    add_equation(doc, var)


def build_document():
    doc = Document()
    configure_document(doc)

    add_paragraph(doc, "דף נוסחאות משולב - גרסה 2", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(doc, "מבוסס על החומר שהבאת, עם ניסוח מסודר ותוספת הסברים קצרים במקומות מועדים לטעויות.", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, italic=True)
    add_note(doc, "מפתח", "הדף מסודר לפי נושאים. ליד נוסחאות מרכזיות הוספתי הערת שימוש קצרה ולא הוכנסו נוסחאות כפולות סתם.")

    add_heading(doc, "1. תוחלת, שונות, קווריאנס והתניה", level=1)
    add_equation(doc, r"E[X]=\sum_x x\,p_X(x)")
    add_equation(doc, r"E[X]=\int_{-\infty}^{\infty} x f_X(x)\,dx")
    add_equation(doc, r"E[aX+b]=aE[X]+b")
    add_equation(doc, r"E\!\left[\sum_i X_i\right]=\sum_i E[X_i]")
    add_equation(doc, r"\mathrm{Var}(X)=E[(X-E[X])^2]=E[X^2]-E[X]^2")
    add_equation(doc, r"\mathrm{Var}(aX)=a^2\mathrm{Var}(X)")
    add_equation(doc, r"\mathrm{Var}(X+Y)=\mathrm{Var}(X)+\mathrm{Var}(Y)+2\mathrm{Cov}(X,Y)")
    add_equation(doc, r"\mathrm{Cov}(X,Y)=E[(X-E[X])(Y-E[Y])]=E[XY]-E[X]E[Y]")
    add_equation(doc, r"\mathrm{Cov}(aX+b,cY+d)=ac\,\mathrm{Cov}(X,Y)")
    add_equation(doc, r"\mathrm{Cov}(X_1+X_2,Y)=\mathrm{Cov}(X_1,Y)+\mathrm{Cov}(X_2,Y)")
    add_equation(doc, r"\mathrm{Cov}(X,X)=\mathrm{Var}(X)")
    add_bullet(doc, 'אם X ו-Y בלתי תלויים אז הקווריאנס 0. הכיוון ההפוך בדרך כלל לא נכון.')
    add_equation(doc, r"E[X]=E[E[X\mid Y]]")
    add_equation(doc, r"\mathrm{Var}(X)=E[\mathrm{Var}(X\mid Y)]+\mathrm{Var}(E[X\mid Y])")
    add_note(doc, "פירוש", "משפט ההחלקה אומר: קודם מחשבים את X כשידוע Y, ואז ממוצעים על כל הערכים של Y.")
    add_equation(doc, r"A\in\{0,1\},\ P(A=1)=p \Longrightarrow E[X]=pE[X\mid A=1]+(1-p)E[X\mid A=0]")
    add_equation(doc, r"P(A,B)=P(B\mid A)P(A)")
    add_equation(doc, r"P(B)=\sum_i P(B\mid A_i)P(A_i)")
    add_equation(doc, r"P(A\mid C)=\sum_k P(A\mid B_k,C)P(B_k\mid C)")
    add_equation(doc, r"P(A\mid B)=\frac{P(A\cap B)}{P(B)}")
    add_equation(doc, r"P(A\mid B)=\frac{P(B\mid A)P(A)}{P(B)}")
    add_equation(doc, r"f_{X\mid Y}(x\mid y)=\frac{f_{Y\mid X}(y\mid x)f_X(x)}{f_Y(y)}")
    add_equation(doc, r"P(A,B,C)=P(A\mid B,C)P(B\mid C)P(C)")
    add_note(doc, "מתי להשתמש", "אם יש משתנה נסתר שמקל את החישוב, מתנים עליו. אם נתונה התוצאה ומחפשים את הסיבה, לרוב בייס הוא הכלי הנכון.")
    add_bullet(doc, "בסכום מקרי או במודל שמחליף בין כמה תרחישים, כמעט תמיד מתחילים מהסתברות שלמה או תוחלת שלמה.")

    add_heading(doc, "2. התפלגויות חשובות, CDF וזהויות עזר", level=1)
    add_distribution_block(
        doc,
        "ברנולי",
        "ניסוי אחד עם הצלחה/כישלון.",
        r"X\sim \mathrm{Bernoulli}(p),\quad P(X=1)=p,\ P(X=0)=1-p",
        r"E[X]=p",
        r"\mathrm{Var}(X)=p(1-p)",
    )
    add_distribution_block(
        doc,
        "בינומית",
        "ספירת מספר הצלחות מתוך n ניסויים בלתי תלויים.",
        r"X\sim \mathrm{Bin}(n,p),\quad P(X=k)=\binom{n}{k}p^k(1-p)^{n-k}",
        r"E[X]=np",
        r"\mathrm{Var}(X)=np(1-p)",
    )
    add_distribution_block(
        doc,
        "פואסון",
        "ספירת אירועים בפרק זמן קבוע.",
        r"X\sim \mathrm{Pois}(\lambda),\quad P(X=k)=e^{-\lambda}\frac{\lambda^k}{k!}",
        r"E[X]=\lambda",
        r"\mathrm{Var}(X)=\lambda",
    )
    add_distribution_block(
        doc,
        "גאומטרית",
        "מספר כישלונות עד הצלחה ראשונה, לפי הסימון בחומר.",
        r"K\sim \mathrm{Geom}(p),\quad P(K=k)=(1-p)^k p,\quad k=0,1,2,\dots",
        r"E[K]=\frac{1-p}{p}",
        r"\mathrm{Var}(K)=\frac{1-p}{p^2}",
    )
    add_distribution_block(
        doc,
        "מעריכית",
        "זמן המתנה בין אירועי פואסון.",
        r"T\sim \mathrm{Exp}(\lambda),\quad f_T(t)=\lambda e^{-\lambda t},\ t\ge 0",
        r"E[T]=\frac{1}{\lambda}",
        r"\mathrm{Var}(T)=\frac{1}{\lambda^2}",
    )
    add_distribution_block(
        doc,
        "נורמלית",
        "גדלים רציפים טבעיים, או תוצאה של סכימת הרבה רכיבים.",
        r"X\sim N(\mu,\sigma^2),\quad f_X(x)=\frac{1}{\sqrt{2\pi\sigma^2}}e^{-\frac{(x-\mu)^2}{2\sigma^2}}",
        r"E[X]=\mu",
        r"\mathrm{Var}(X)=\sigma^2",
    )
    add_distribution_block(
        doc,
        "אחידה",
        "לכל ערך בטווח יש אותו סיכוי.",
        r"X\sim U[a,b],\quad f_X(x)=\frac{1}{b-a},\ a\le x\le b",
        r"E[X]=\frac{a+b}{2}",
        r"\mathrm{Var}(X)=\frac{(b-a)^2}{12}",
    )
    add_equation(doc, r"F_X(\alpha)=P(X\le \alpha)")
    add_equation(doc, r"\Phi(-\alpha)=1-\Phi(\alpha)")
    add_equation(doc, r"P(a\le X\le b)=\Phi(b)-\Phi(a)")
    add_equation(doc, r"X\sim N(0,1) \Longrightarrow E[|X|]=\sqrt{\frac{2}{\pi}}")
    add_equation(doc, r"\phi_X(t)=E[e^{itX}]=\exp\left(it\mu-\frac{1}{2}\sigma^2 t^2\right)")
    add_equation(doc, r"X\sim N(0,\sigma^2) \Longrightarrow E[X^4]=3\sigma^4")
    add_bullet(doc, "אם יש ערך מוחלט של משתנה גאוסי, בדרך כלל המשתנה החדש כבר לא גאוסי.")
    add_equation(doc, r"e^{j\theta}=\cos\theta+j\sin\theta")
    add_equation(doc, r"\cos(A)\cos(B)=\frac{1}{2}\bigl(\cos(A-B)+\cos(A+B)\bigr)")
    add_equation(doc, r"\cos(2\pi f_0 t)=\frac{e^{j2\pi f_0 t}+e^{-j2\pi f_0 t}}{2}")
    add_equation(doc, r"\cos^2(x)=\frac{1+\cos(2x)}{2}")
    add_equation(doc, r"\operatorname{sinc}(x)=\frac{\sin(\pi x)}{\pi x}")
    add_equation(doc, r"g(t)\cos(2\pi f_c t)\ \Longleftrightarrow\ \frac{1}{2}\bigl(G(f-f_c)+G(f+f_c)\bigr)")
    add_equation(doc, r"S_{x\cos}(f)=\frac{1}{4}\bigl(S_X(f-f_0)+S_X(f+f_0)\bigr)")
    add_bullet(doc, "כפל בזמן ב-cos מזיז ספקטרום ימינה ושמאלה; לכן ב-PSD מתקבלים שני עותקים מוקטנים פי 1/4.")
    add_equation(doc, r"\sum_{j=0}^{M}\binom{M}{j}=2^M")
    add_equation(doc, r"\begin{pmatrix}a&b\\c&d\end{pmatrix}^{-1}=\frac{1}{ad-bc}\begin{pmatrix}d&-b\\-c&a\end{pmatrix}")

    add_heading(doc, "3. וקטור גאוסי וגאוסיות במשותף", level=1)
    add_equation(doc, r"X=(X_1,\dots,X_n)^T\ \mathrm{jointly\ Gaussian}\iff a^T X\ \mathrm{Gaussian\ for\ every}\ a")
    add_equation(doc, r"\phi_X(t)=E[e^{it^T X}]=\exp\left(it^T\mu-\frac{1}{2}t^T\Sigma t\right)")
    add_equation(doc, r"Y=AX+b,\ X\sim N(\mu,\Sigma)\Longrightarrow Y\sim N(A\mu+b,A\Sigma A^T)")
    add_bullet(doc, "כל תת-וקטור של וקטור גאוסי הוא גם גאוסי במשותף.")
    add_bullet(doc, "אם כל רכיב גאוסי לחוד, זה עדיין לא מספיק כדי להסיק גאוסיות במשותף.")
    add_equation(doc, r"\mathrm{Cov}(U,V)=0,\ (U,V)\ \mathrm{jointly\ Gaussian}\Longrightarrow U\perp V")
    add_bullet(doc, "באופן כללי: אי-תלות גוררת חוסר קורלציה, אבל חוסר קורלציה גורר אי-תלות רק במקרה הגאוסי המשותף.")
    add_equation(doc, r"E[X\mid Y]=E[X]+\frac{\mathrm{Cov}(X,Y)}{\mathrm{Var}(Y)}(Y-E[Y])")
    add_equation(doc, r"\mathrm{Var}(X_1\mid X_2)=\mathrm{Var}(X_1)-\frac{\mathrm{Cov}(X_1,X_2)^2}{\mathrm{Var}(X_2)}")
    add_note(doc, "שימו לב", "הנוסחה האחרונה נכונה כתוחלת מותנית ליניארית/אופטימלית, ובמקרה הגאוסי המשותף היא גם ה-MMSE האמיתי.")
    add_equation(doc, r"E[X\mid Y]=E[X]+\Sigma_{XY}\Sigma_Y^{-1}(Y-E[Y])")
    add_equation(doc, r"\Sigma_{XY}=\begin{pmatrix}\mathrm{Cov}(X,Y_1)&\cdots&\mathrm{Cov}(X,Y_n)\end{pmatrix}")
    add_equation(doc, r"\mathrm{Var}(X_1\mid X_2,X_3)=s_{11}-\begin{pmatrix}s_{12}&s_{13}\end{pmatrix}\begin{pmatrix}s_{22}&s_{23}\\s_{23}&s_{33}\end{pmatrix}^{-1}\begin{pmatrix}s_{12}\\s_{13}\end{pmatrix}")
    add_bullet(doc, "אם תופיע פונקציה לא ליניארית כמו sign(X), אל תניח אוטומטית גאוסיות במשותף.")

    add_heading(doc, "4. תהליך פואסון וזמני המתנה", level=1)
    add_equation(doc, r"N(t_1)-N(t_0)\sim \mathrm{Pois}\bigl(\lambda(t_1-t_0)\bigr)")
    add_equation(doc, r"P(N(t)=k)=e^{-\lambda t}\frac{(\lambda t)^k}{k!}")
    add_equation(doc, r"E[N(t)]=\lambda t,\qquad \mathrm{Var}(N(t))=\lambda t")
    add_equation(doc, r"E[N(a)N(b)]=\lambda^2ab+\lambda\min(a,b)")
    add_bullet(doc, "בקטעים לא חופפים, מספרי ההגעות בלתי תלויים.")
    add_equation(doc, r"T_1\sim \mathrm{Exp}(\lambda)")
    add_equation(doc, r"P(T_1>s)=e^{-\lambda s}")
    add_equation(doc, r"E[T_i-T_{i-1}]=\frac{1}{\lambda}")
    add_note(doc, "מתי להתנות", "אם הקצב עצמו אקראי, או אם N מופיע בתוך משתנה אחר, מתנים קודם על המשתנה שקובע את הקצב ואז מחליקים.")
    add_equation(doc, r"E[N(X)]=E[E[N(X)\mid X]]=E[\lambda X]")
    add_bullet(doc, "אם תערובת של תהליכי פואסון לא מחזירה התפלגות פואסונית, זה לא תהליך פואסון.")
    add_bullet(doc, "מכפלה של שני תהליכי פואסון בדרך כלל שוברת את התנאי של לכל היותר קפיצה אחת קטנה בזמן קצר.")

    add_heading(doc, "5. סטציונריות, סמ\"ר, אוטוקורלציה ו-PSD", level=1)
    add_equation(doc, r"F_{X(t_1+\tau),\dots,X(t_n+\tau)}=F_{X(t_1),\dots,X(t_n)}")
    add_equation(doc, r"\mu_X[n]=E[X[n]]=\mu")
    add_equation(doc, r"R_X[n,n+k]=E[X[n]X[n+k]]=r_X[k]")
    add_equation(doc, r"R_{XY}[k]=E[X[n]Y[n+k]]")
    add_equation(doc, r"R_X(t_1,t_2)=E[X_{t_1}X_{t_2}]")
    add_equation(doc, r"R_X(\tau)=E[X_tX_{t+\tau}]")
    add_equation(doc, r"R_{XY}(t_1,t_2)=E[X_{t_1}Y_{t_2}]")
    add_equation(doc, r"R_X(\tau)=R_X(-\tau)")
    add_equation(doc, r"|R_X(\tau)|\le R_X(0)")
    add_equation(doc, r"E[(X_{t+\tau}-X_t)^2]=2(R_X(0)-R_X(\tau))")
    add_note(doc, "פירוש", "סטציונרי חזק שומר על כל ההתפלגות תחת הזזה בזמן. סמ\"ר שומר רק על התוחלת ועל הקורלציה.")
    add_equation(doc, r"R_X[k]=\sigma^2\delta[k]")
    add_equation(doc, r"S_X(f)=\mathcal{F}\{R_X(\tau)\}")
    add_equation(doc, r"R_X(0)=\int_{-\infty}^{\infty}S_X(f)\,df")
    add_equation(doc, r"E[X^2(t)]=\mu_X^2+\int_{-\infty}^{\infty}S_X(f)\,df")
    add_equation(doc, r"Y=h*X,\ X\ \mathrm{WSS}\Longrightarrow \mu_Y=H(0)\mu_X")
    add_equation(doc, r"R_Y(\tau)=(R_X*h*\tilde h)(\tau)")
    add_equation(doc, r"S_Y(f)=|H(f)|^2S_X(f)")
    add_equation(doc, r"R_{XY}(\tau)=(R_X*h)(\tau)")
    add_equation(doc, r"S_{XY}(f)=S_X(f)H(f)")
    add_equation(doc, r"Y_1=h_1*X,\ Y_2=h_2*X \Longrightarrow S_{Y_1Y_2}(f)=S_X(f)H_1(f)\overline{H_2(f)}")
    add_equation(doc, r"E[Y^2(t)]=R_Y(0)=\int_{-\infty}^{\infty}S_X(f)|H(f)|^2\,df")
    add_bullet(doc, "אם תהליך עובר דרך מערכת LTI יציבה, והקלט סמ\"ר, גם הפלט סמ\"ר.")
    add_bullet(doc, "תהליך עם תנאי עצירה, שינוי חוקים בזמן, או התחלה מיוחדת הוא חשוד מיידית כלא-סטציונרי.")

    add_heading(doc, "6. שערוך MMSE, שערוך ליניארי ומסנן וינר", level=1)
    add_equation(doc, r"\hat X_{\mathrm{MMSE}}=E[X\mid Y]")
    add_equation(doc, r"\hat X=aY+b,\qquad a=\frac{\mathrm{Cov}(X,Y)}{\mathrm{Var}(Y)},\qquad b=E[X]-aE[Y]")
    add_equation(doc, r"\mathrm{MSE}=E[(X-\hat X)^2]")
    add_note(doc, "פירוש", "MMSE הוא האומדן הכי טוב ללא אילוץ על הצורה. אם מגבילים את האומדן להיות ליניארי, פותרים עבור a,b.")
    add_equation(doc, r"H_{\mathrm{opt}}(f)=\frac{S_{XY}(f)}{S_Y(f)}")
    add_equation(doc, r"S_e(f)=S_X(f)-\frac{|S_{XY}(f)|^2}{S_Y(f)}")
    add_equation(doc, r"E[e^2]=\int_{-\infty}^{\infty}S_e(f)\,df")
    add_equation(doc, r"Y=X+N,\ X\perp N\Longrightarrow H_{\mathrm{opt}}(f)=\frac{S_X(f)}{S_X(f)+S_N(f)}")
    add_equation(doc, r"E[e^2]=\int_{-\infty}^{\infty}\Bigl(S_X(f)|1-H(f)|^2+S_N(f)|H(f)|^2\Bigr)\,df")
    add_bullet(doc, "בפתרון מלא של וינר: מוצאים קודם אוטוקורלציה/קרוס-קורלציה, מעבירים ל-PSD, ורק אז מציבים ב-H_opt.")
    add_bullet(doc, "אם S_XY(f)=0 לכל f, המסנן האופטימלי הוא 0 והאומדן הטוב ביותר הוא הקבוע E[X].")
    add_bullet(doc, "במקרה של גאוסיות משותפת, MMSE ו-LMMSE מתלכדים.")

    add_heading(doc, "7. שרשראות מרקוב", level=1)
    add_equation(doc, r"P(X_{n+1}=j\mid X_n=i,X_{n-1},\dots,X_0)=P(X_{n+1}=j\mid X_n=i)")
    add_equation(doc, r"P=[p_{ij}]")
    add_equation(doc, r"v(n)=v(0)P^n")
    add_equation(doc, r"P(X_0=X_1)=\sum_i v_i(0)p_{ii}")
    add_equation(doc, r"P(X_0=i\mid X_n=j)=\frac{(P^n)_{ij}v_i(0)}{\sum_m (P^n)_{mj}v_m(0)}")
    add_equation(doc, r"P(X_{2n}=k\mid X_n=j)=(P^n)_{jk}")
    add_equation(doc, r"\nu P=\nu,\qquad \sum_j \nu_j=1")
    add_bullet(doc, "הומוגניות היא תכונה של חוקי המעבר. סטציונריות היא תכונה של התפלגות המצבים.")
    add_bullet(doc, "מחלקה קשירה וסגורה נותנת מצבים נשנים; מחלקה פתוחה נותנת מצבים חולפים.")
    add_bullet(doc, "פילוג סטציונרי יכול להיות שונה מאפס רק על מצבים נשנים.")
    add_equation(doc, r"E_j[N_i]=\sum_{n=0}^{\infty}(P^n)_{ji}")
    add_equation(doc, r"E[N_i\mid X(0)=j]=\frac{\rho_{ji}}{1-\rho_{ii}}\qquad (\rho_{ii}<1)")
    add_note(doc, "מתי מזהים מרקוב", "אם אפשר לכתוב x_n=f(x_{n-1},V_n) כאשר הרעש החדש בלתי תלוי בעבר, בדרך כלל יש מבנה מרקובי. אם גם V_n הוא i.i.d, נקבל לרוב הומוגניות.")

    add_heading(doc, "8. תבניות פתרון מתוך החומר", level=1)
    add_bullet(doc, "כש-Y נבנה ממשתנה בוחר A, למשל Y=AX+(1-A)N, מתנים קודם על A ואז מחזירים ל-Y עם בייס או החלקה.")
    add_equation(doc, r"E[X\mid Y]=E[E[X\mid A,Y]\mid Y]")
    add_bullet(doc, "כשרוצים לבדוק אם A ו-Y בלתי תלויים, בודקים אם הפילוג המותנה של Y בהינתן A זהה לפילוג השולי של Y.")
    add_bullet(doc, "כשמופיעים sign(X) או |X|, נסו לעבור חזרה ל-X דרך סימטריה או חלוקה למקרים.")
    add_bullet(doc, "כדי לבדוק אם תהליך הוא פואסון, לא מספיק לחשב תוחלת. צריך שההתפלגות של N(t) תהיה מהצורה הפואסונית וגם שהתוספות בקטעים זרים יהיו בלתי תלויות.")
    add_bullet(doc, "כדי לבדוק סטציונריות של תהליך עם אינדקס אקראי, משתמשים בהסתברות שלמה על האינדקס ואז מנצלים בת\"ס/סימטריה אם קיימת.")
    add_bullet(doc, "במרקוב, כשמבקשים תוחלת זמן פגיעה או מספר ביקורים, התניה על הצעד הראשון היא בדרך כלל המהלך הנכון.")

    add_heading(doc, "9. מלכודות נפוצות", level=1)
    add_bullet(doc, "חוסר קורלציה אינו אי-תלות, אלא אם יש גאוסיות במשותף.")
    add_bullet(doc, "בתלות בזוגות אינה מספיקה לאי-תלות של הווקטור כולו.")
    add_bullet(doc, "MMSE אינו בהכרח ליניארי; הוא נהיה ליניארי במקרה הגאוסי המשותף או כשמגבילים מראש לצורה ליניארית.")
    add_bullet(doc, "סמ\"ר אינו זהה לסטציונריות חזקה.")
    add_bullet(doc, "במרקוב: p_{ii}=1 אינו ההגדרה לנשנות. נשנות קשורה לחזרה עתידית בהסתברות 1, לא רק להישארות בצעד אחד.")
    add_bullet(doc, "אם X ו-Y בלתי תלויים אז Cov=0. להשתמש בכיוון ההפוך רק אם יש סיבה טובה, בדרך כלל גאוסיות במשותף.")

    add_paragraph(doc, "הערה אחרונה: כשיש ספק, חוזרים להגדרה. במבחן זה כמעט תמיד יותר בטוח מהפעלת נוסחה לא בתנאים שלה.", bold=True, space_after=0)
    doc.save(str(OUTPUT))


if __name__ == "__main__":
    build_document()
    print(OUTPUT.as_posix().encode("ascii", "backslashreplace").decode("ascii"))
